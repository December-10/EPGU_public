from docx import Document
import pyperclip
import logging
import pyautogui
import time
import keyboard
import re
from datetime import datetime, timedelta
import fnmatch
import os
import shutil                          
import subprocess                      
from pathlib import Path              
from openpyxl import load_workbook      
from PyPDF2 import PdfReader        
import sys    
from types import NoneType
from PyQt5.QtWidgets import (
    QApplication, QWidget, QLineEdit, QComboBox,
    QPushButton, QCheckBox, QFileDialog, QHBoxLayout,
    QVBoxLayout, QMessageBox)
from PyQt5.QtGui import QPalette, QColor, QCloseEvent
from PyQt5.QtCore import Qt

if __name__ == '__main__':

    def extract_data(data_name, expert_organization=None, year=2025):
        """
        Извлекает данные из файла "Исходные данные (публ).docx"
        """

        docx_doc = Document('_internal/Публикуемое/Исходные данные (публ).docx')

        year_table_dict = {2024:0,
                           2025:1}

        table_number = year_table_dict[int(year)]

        for row in docx_doc.tables[table_number].rows:
            if data_name in row.cells[0].text:
                if expert_organization is None or expert_organization == "Экспертная организация 1":
                    return row.cells[1].text
                elif expert_organization == "Экспертная организация 2":
                    return row.cells[2].text

    # задаем постоянные данные
    customer_region = extract_data("Регион")
    telephone = extract_data("Телефон")
    e_mail = extract_data("E-Mail")
    confidence = float(extract_data("Точность поиска изображений"))
    downloads_directory = extract_data("Директория загрузок")
    signing_statement_confidence = float(extract_data("Точность поиска области подписания заявления"))
    startpoint_x_deviation = int(extract_data("Отклонение от области подписания по оси X в стартовую точку"))
    startpoint_y_deviation = int(extract_data("Отклонение от области подписания по оси Y в стартовую точку"))
    endpoint_x_deviation = int(extract_data("Отклонение от области подписания по оси X в конечную точку"))
    endpoint_y_deviation = int(extract_data("Отклонение от области подписания по оси Y в конечную точку"))
    signing_speed = float(extract_data("Скорость подписания заявления с помощью мыши"))
    signing_confidence_reduction_speed = float(extract_data("Скорость снижения точности поиска изображений при подписании"))
    acrobat_reader_path = extract_data("Путь к программе Acrobat Reader")
    opo_address_directory = extract_data("Директория расположения файла с адресами ОПО")
    opo_address_file = Document(opo_address_directory)

    def check_file_open(check_file_name):
        """
        Проверяет, открыт ли файл и возвращает True(открыт) или False(закрыт)
        """
        path_to_file = check_file_name.split('\\') if '\\' in check_file_name else check_file_name.split('/')
        directory = '\\'.join(path_to_file[:-1])
        file_name = path_to_file[-1]
        return True if os.path.isfile(Path(directory, f'~${file_name}')) is True else False

    def paste_info(name):
        """
        Вставляет переданную в функцию информацию, предварительно скопировав её в буфер обмена
        """
        pyperclip.copy(name)
        pyperclip.paste()
        keyboard.press_and_release("ctrl+v")

    def click_on_pic (image, confidence=confidence, x_offset=0, y_offset=0, action='click'):
        """
        Нажимает в центр переданного в функцию и найденного на экране изображения
        """

        dir = '/'.join(image.split('/')[:-1])    # извлекаем директорию расположения файла
        pattern = image.split('/')[-1].replace('.png', '*')    # извлекаем имя файла (расширение заменяем на звездочку)

        click_flag = 0
        while click_flag == 0:    # прогоняем цикл, пока не нажмем на изображение
            list = []
            for root, dirs, files in os.walk(dir):
                for name in files:
                    if fnmatch.fnmatch(name, pattern):
                        list.append(root + '/' + name)
            # к искомым изображениям также добавляем изображение окна с ошибкой
            list.append(root + '/' + '81_Oshibka.png')

            for image_file in list:
                try:
                    location = pyautogui.locateOnScreen(image_file, confidence=confidence)
                    point = pyautogui.center(location)
                    x, y = point
                    x += x_offset
                    y += y_offset
                    if action == 'click':
                        pyautogui.click(x, y)
                    elif action == 'moveTo':
                        pyautogui.moveTo(x, y)
                    click_flag = 1
                    break  # выход из цикла после успешного клика
                except pyautogui.ImageNotFoundException:
                    continue

    # создаем словарь для замен значений в будущем
    date_dict = {' января ': '.01.',
                 ' февраля ': '.02.',
                 ' марта ': '.03.',
                 ' апреля ': '.04.',
                 ' мая ': '.05.',
                 ' июня ': '.06.',
                 ' июля ': '.07.',
                 ' августа ': '.08.',
                 ' сентября ': '.09.',
                 ' октября ': '.10.',
                 ' ноября ': '.11.',
                 ' декабря ': '.12.'}

    # настраиваем логирование:
    logging.basicConfig(level=logging.DEBUG, filename="ZEPB.log", force=True, encoding='utf-8',
                        format="%(asctime)s (%(levelname)s) %(message)s", datefmt='%d.%m.%Y %H:%M:%S')

    # настраиваем окно ввода данных
    class MyForm(QWidget):
        def __init__(self):
            super().__init__()
            self.setWindowTitle("Ввод данных")  # заголовок
            self.init_ui()  # функция формирования окна

            # заполняемые формы
            self.selected_files = ""
            self.zepb_number = ""
            self.selected_city = ""
            self.selected_org = ""
            self.selected_year = ""
            self.only_check = False

        def init_ui(self):
            layout = QVBoxLayout()

            zepb_number_input_layout = QHBoxLayout()
            self.zepb_number_text = QLineEdit()
            self.zepb_number_text.setPlaceholderText("Введите номер(-а) заключения(-й)")
            zepb_number_input_layout.addWidget(self.zepb_number_text)
            self.file_button = QPushButton("Выбрать заявление")
            self.file_button.clicked.connect(self.choose_file)
            zepb_number_input_layout.addWidget(self.file_button)
            layout.addLayout(zepb_number_input_layout)

            self.year_combo = QComboBox()
            self.year_combo.addItems(['2024', '2025'])
            layout.addWidget(self.year_combo)

            self.city_combo = QComboBox()
            self.city_combo.addItems(["Красноярск", "Томск", "Новосибирск"])
            layout.addWidget(self.city_combo)

            self.org_combo = QComboBox()
            self.org_combo.addItems(["Экспертная организация 1", "Экспертная организация 2"])
            layout.addWidget(self.org_combo)

            self.checkbox = QCheckBox("Только проверка заявления")
            layout.addWidget(self.checkbox)

            button_layout = QHBoxLayout()
            self.ok_button = QPushButton("Ок")
            self.ok_button.clicked.connect(self.on_ok_clicked)
            self.cancel_button = QPushButton("Отмена")
            self.cancel_button.clicked.connect(self.close)
            button_layout.addWidget(self.ok_button)
            button_layout.addWidget(self.cancel_button)
            layout.addLayout(button_layout)

            self.setLayout(layout)

        def choose_file(self):
            files, _ = QFileDialog.getOpenFileNames(
                self,
                "Выберите заявления",
                "",
                "Все файлы (*.*)"
            )
            if files:
                self.selected_files = files  # сохраняем список файлов
                self.file_button.setText(f"Файлов выбрано: {len(files)}")  # обновляем текст на кнопке
                self.zepb_number_text.setText(";".join(files))  # ставим список выбранных файлов в полее ввода

        def on_ok_clicked(self):
            self.close()  # Закрываем форму

        def closeEvent(self, event: QCloseEvent):
            exit()

    def message(text):
        msg = QMessageBox(form)
        msg.setWindowTitle("Информация")
        msg.setText(text)
        msg.exec_()

    # читаем окно
    app = QApplication(sys.argv)
    form = MyForm()
    form.resize(350, 300)
    # ставим темную тему
    dark_palette = QPalette()
    dark_palette.setColor(QPalette.Window, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.WindowText, Qt.white)
    dark_palette.setColor(QPalette.Base, QColor(25, 25, 25))
    dark_palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.ToolTipBase, Qt.white)
    dark_palette.setColor(QPalette.ToolTipText, Qt.white)
    dark_palette.setColor(QPalette.Text, Qt.white)
    dark_palette.setColor(QPalette.Button, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.ButtonText, Qt.white)
    dark_palette.setColor(QPalette.BrightText, Qt.red)
    dark_palette.setColor(QPalette.Link, QColor(42, 130, 218))
    dark_palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
    dark_palette.setColor(QPalette.HighlightedText, Qt.black)
    app.setPalette(dark_palette)
    app.setStyle("Fusion")

    form.show()
    app.exec_()  # Запуск GUI и ожидание закрытия

    # выводим данные в переменные
    input = form.zepb_number_text.text()
    rtn_region = form.city_combo.currentText()
    exp_organization = form.org_combo.currentText()
    year = form.year_combo.currentText()
    only_check = form.checkbox.isChecked()

    # если вводом являлся список файлов, то разделитель принимаем ";", в ином случае - " "(пробел)
    separator = ';' if form.selected_files else ' '

    processed_files_list = []    # список для внесения успешно обработанных заключений
    error_files_list = []    # список для внесения не обработанных заключений

    # задаем оставшиеся индивидуальные данные
    statement_base_number = extract_data("Базовый номер заявления", expert_organization=exp_organization, year=year)
    zepb_registration_number_key = extract_data("Ключевое значение для поиска регистрационного номера заключения", expert_organization=exp_organization, year=year)
    zepb_registration_number_lenght = int(extract_data("Длина регистрационного номера заключения", expert_organization=exp_organization, year=year))
    zepb_sheet_name = str(extract_data("Наименование листа файла контроля с данными по заключениям", expert_organization=exp_organization, year=year))
    zepb_directory = extract_data("Директория расположения заключений ЭПБ", expert_organization=exp_organization, year=year)
    license_number = extract_data("Номер лицензии", expert_organization=exp_organization, year=year)
    statement_directory = extract_data("Директория расположения заявлений", expert_organization=exp_organization, year=year)
    consolidated_file_directory = extract_data("Директория расположения файла с контролем", expert_organization=exp_organization, year=year)

    # определяем, не открыт ли файл сводной и если открыт, то программу не запускаем:
    if check_file_open(consolidated_file_directory) is True and only_check is not True:
        message(f'Файл {consolidated_file_directory} открыт.\nЗакройте его и запустите программу повторно.')
        exit()

    # вносим данные с файла сводной в словарь
    while True:
        try:
            wb_cf = load_workbook(consolidated_file_directory, data_only=False)
            break
        except:
            message(f'Не удалось открыть файл {consolidated_file_directory}.'
                           'Проверьте его, после чего нажмите Ок')
            continue

    try:
        ws_cf = wb_cf[zepb_sheet_name]
    except:
        ws_cf = wb_cf.active

    min_row = ws_cf.min_row    # номер первой заполненной строки
    max_row = ws_cf.max_row    # номер последней заполненной строки

    # номер столбца, в котором содержатся полные номера заключений
    column = None
    for cells in ws_cf['A1:Z10']:
        for cell in cells:
            if cell.value == 'Номер ЗЭПБ':
                column = cell.column
                break

    min_cell = ws_cf.cell(min_row, column).coordinate    # адрес первой заполненной ячейки
    max_cell = ws_cf.cell(max_row, column).coordinate    # адрес последней заполненной ячейки

    consolidated_file_dict = {}

    for cell in ws_cf[f'{min_cell}:{max_cell}']:
        consolidated_file_dict[cell[0].value] = {'coordinate':cell[0].coordinate,
                                            'status':cell[0].offset(column=3).value}

    # запускаем обработку заявлений
    for path_to_statement_file in input.split(separator):
        try:

            exceptions_list = []

            file_check = path_to_statement_file    # маркер для проверки, является ли текущий файл первым введенным

            # на примере номера заключения "028/04-24-ТУ/НГ/1246/057":
            if exp_organization == 'Экспертная организация 1':
                # изменяем путь к файлу в зависимости от введенного значения
                if len(path_to_statement_file) < 3:    # если передан только ключевой номер заявления для зданий (например, "2")
                    path_to_statement_file = fr"{statement_directory}\{path_to_statement_file}.docx"
                if len(path_to_statement_file) == 3:    # если передан только ключевой номер заявления (например, "057")
                    path_to_statement_file = fr"{statement_directory}\{statement_base_number}-{path_to_statement_file}.docx"
                # если передан базовый и ключевой номер заявления через тире (например, "1246-057")
                elif len(path_to_statement_file) == 8 and "-" in path_to_statement_file:
                    path_to_statement_file = fr"{statement_directory}\{path_to_statement_file}.docx"
                # если передан базовый и ключевой номер заявления через слэш (например, "1246/057")
                elif len(path_to_statement_file) == 8 and "/" in path_to_statement_file:
                    path_to_statement_file = fr"{statement_directory}\{path_to_statement_file.replace('/', '-')}.docx"
                # если передан полный номер заявления (например, "028/04-24-ТУ/НГ/1246/057")
                elif len(path_to_statement_file) == 24:
                    path_to_statement_file = fr"{statement_directory}\{path_to_statement_file[16:].replace('/', '-')}.docx"
                # если передан иной номер заявления (например, "24008-2/9")
                elif len(path_to_statement_file) >= 9 and statement_base_number not in path_to_statement_file:
                    zepb_registration_number_key = path_to_statement_file[:5]
                    zepb_registration_number_lenght = len(path_to_statement_file)
                    path_to_statement_file = fr"{statement_directory}\{path_to_statement_file.replace('/', '_')}.docx"

                # извлекаем имя файла:
                for i, symbol in enumerate(str(path_to_statement_file)[::-1]):
                    if symbol == '/' or symbol == '\\':
                        if statement_base_number in path_to_statement_file:
                            file_name = str(path_to_statement_file)[-i:str(path_to_statement_file).find('.docx')].replace("_", "-")
                        else:
                            file_name = str(path_to_statement_file)[-i:str(path_to_statement_file).find('.docx')]
                        file_name = file_name.replace("Заявление", "")
                        file_name = file_name.replace("№", "")
                        file_name = file_name.replace(" ", "")
                        break

                path_to_zepb = (rf"{zepb_directory}\{file_name}.pdf")


            elif exp_organization == 'Экспертная организация 2':
                # на примере номера заключения "001-2024-028/009":
                # изменяем путь к файлу в зависимости от введенного значения
                if len(path_to_statement_file) == 3 or len(path_to_statement_file) == 4:  # если передан только ключевой номер заявления (например, "009" или "1250")
                    path_to_statement_file = fr"{statement_directory}\001-{year}-{statement_base_number}_{path_to_statement_file}.xlsx"
                # если передан базовый и ключевой номер заявления через тире (например, "028-009")
                elif len(path_to_statement_file) == 7 and "-" in path_to_statement_file:
                    path_to_statement_file = fr"{statement_directory}\001-{year}-{path_to_statement_file.replace('-', '_')}.xlsx"
                # если передан базовый и ключевой номер заявления через слэш (например, "028/009")
                elif len(path_to_statement_file) == 7 and "/" in path_to_statement_file:
                    path_to_statement_file = fr"{statement_directory}\001-{year}-{path_to_statement_file.replace('/', '_')}.xlsx"
                # если передан полный номер заявления (например, "001-2024-028_009" или "001-2024-028/009")
                elif len(path_to_statement_file) == 16:
                    path_to_statement_file = fr"{statement_directory}\{path_to_statement_file.replace('/', '_')}.xlsx"

                # извлекаем имя файла:
                for i, symbol in enumerate(str(path_to_statement_file)[::-1]):
                    if symbol == '/' or symbol == '\\':
                        file_name = str(path_to_statement_file)[-i:str(path_to_statement_file).find('.xlsx')]
                        file_name = file_name.replace(f"001-{year}-", "")
                        file_name = file_name.replace(" ", "")
                        break

                path_to_zepb = (rf"{zepb_directory}\001-{year}-{file_name}.pdf")

            if os.path.isfile(path_to_statement_file) is False:
                exceptions_list.append(f'Файл заявления отсутствует ({path_to_statement_file})')

            if os.path.isfile(path_to_zepb) is False:
                exceptions_list.append(f'Файл заключения отсутствует ({path_to_zepb})')

            if exp_organization == 'Экспертная организация 1':
                # открываем заявление в формате docx
                document = Document(path_to_statement_file)
                paragraphs = document.paragraphs
                tables = document.tables
                table1 = tables[0]  # Адресат заявления
                table2 = tables[1]  # Сведения о заявителе: Юридическое лицо
                table3 = tables[2]  # Индивидуальный предприниматель
                table4 = tables[3]  # Сведения об экспертной организации
                table5 = tables[4]  # ФИО и № Эксперта
                table6 = tables[5]  # Наименование заключения ЭПБ
                table7 = tables[6]  # Краткая характеристика объекта экспертизы
                table8 = tables[7]  # Сведения об эксплуатирующей организации

                # задаем данные по заявлению
                expert_organization_full_name = table4.rows[0].cells[1].text.strip()
                expert_organization_INN = table4.rows[2].cells[1].text.strip()
                expert_organization_OGRN = table4.rows[3].cells[1].text.strip()
                zepb_name = table6.rows[0].cells[0].text.strip().replace('\n', ' ')
                zepb_object_name_full_row = table7.rows[0].cells[1].text.strip()
                identification_number = table7.rows[1].cells[1].text
                customer_full_name = table2.rows[0].cells[1].text.strip()
                customer_INN = table2.rows[2].cells[1].text.strip()
                assignment_index = zepb_object_name_full_row.find('предназначен')
                zepb_object_name = zepb_object_name_full_row[:assignment_index]
                zepb_object_assignment = zepb_object_name_full_row[assignment_index:]
                opo_number = table8.rows[6].cells[1].text.replace('№', '').strip()
                # извлекаем квалификационные номера экспертов:
                table5_text = []
                for row in table5.rows:
                    for cell in row.cells:
                        table5_text.append(cell.text)
                table5_text = ' '.join(table5_text)
                expert_number_list = list(set(re.findall(r'\D\D\.\d\d\.\d\d\d\d\d\.\d\d\d', table5_text)))
                # цикл для поиска параграфа с регистрационным номером заключения ЭПБ:
                n = 0
                while n <= len(paragraphs):
                    if "Регистрационный номер заключения экспертизы промышленной безопасности" in paragraphs[n].text:
                        zepb_reg_num_full_row = paragraphs[n].text
                        break
                    n += 1
                zepb_reg_num_idx_start = zepb_reg_num_full_row.find(zepb_registration_number_key)
                zepb_reg_num_idx_end = zepb_reg_num_idx_start + zepb_registration_number_lenght
                zepb_reg_num = zepb_reg_num_full_row[zepb_reg_num_idx_start:zepb_reg_num_idx_end]
                if ' ' in zepb_reg_num:
                    zepb_reg_num = zepb_reg_num.replace(' ', '-').replace('№', '')
                # цикл для поиска параграфа с датой подписания заключения ЭПБ:
                n = 0
                while n <= len(paragraphs):
                    if "Дата подписания заключения" in paragraphs[n].text:
                        zepb_sign_date_full_row = paragraphs[n].text
                        break
                    n += 1
                # цикл для поиска параграфа со сроком дальнейшей безопасной эксплуатации:
                n = 0
                while n <= len(paragraphs):
                    if "Срок дальнейшей безопасной эксплуатации" in paragraphs[n].text:
                        zepb_exploperiod_full_row = paragraphs[n + 1].text  # берем срок безопасной эксплуатации из следующего параграфа
                        break
                    n += 1

            elif exp_organization == 'Экспертная организация 2':
                wb = load_workbook(path_to_statement_file, data_only=False)
                ws = wb.active

                if 'удостоверения' in ws['A7'].value:
                    raise Exception(f'Структура заявления нарушена (обратить внимание на строку 7)')

                empty_cells = []
                for cell in ['B2', 'B3', 'B4', 'B5', 'B6', 'B7', 'B8', 'B9', 'B10', 'B13', 'B14', 'B17', 'B20', 'B21', 'B22']:
                    if type(ws[cell].value) is NoneType:
                        empty_cells.append(cell)

                if empty_cells != []:
                    if len(empty_cells) == 1:
                        exceptions_list.append(f'Ячейка {', '.join(empty_cells)} в заявлении пуста')
                    else:
                        exceptions_list.append(f'Ячейки {', '.join(empty_cells)} в заявлении пусты')

                expert_organization_full_name = ws['B2'].value
                expert_organization_INN = ws['B3'].value
                expert_organization_OGRN = ws['B4'].value
                expert_number_list = list(set(re.findall(r'\D\D\.\d\d\.\d\d\d\d\d\.\d\d\d', ws['B6'].value)))
                zepb_name = ws['B7'].value.replace('\n', ' ').replace('  ', ' ')
                zepb_object_name_full_row = ws['B9'].value.replace('\n', ' ').replace('  ', ' ')

                # убираем предлог "на" из наименования объекта экспертизы:
                while (zepb_object_name_full_row[0] == 'н' or
                       zepb_object_name_full_row[0] == 'а' or
                       zepb_object_name_full_row[0] == ' '):
                    zepb_object_name_full_row = zepb_object_name_full_row[1:]
                zepb_object_name = zepb_object_name_full_row
                zepb_object_assignment = ws['B10'].value.replace('\n', ' ').replace('  ', ' ')
                identification_number = ws['B12'].value if ws['B12'].value is not None else ws['B11'].value

                if identification_number is None:
                    exceptions_list.append('Идентификационный номер не внесен в заявление')
                identification_number = (identification_number
                                         .replace('\n', ' ')
                                         .replace('  ', ' ')
                                         .replace('. №', '.№'))
                customer_full_name = ws['B13'].value.strip()
                customer_INN = ws['B14'].value
                opo_number = ws['B17'].value.strip()
                zepb_reg_num_full_row = ws['B20'].value.strip()
                zepb_reg_num = zepb_reg_num_full_row.replace('№', '').replace(' ', '')
                zepb_sign_date_full_row = ws['B21'].value
                zepb_exploperiod_full_row = ws['B22'].value

                wb.close()

            if zepb_reg_num.replace('№', '').replace(' ', '') in consolidated_file_dict.keys():
                if consolidated_file_dict[zepb_reg_num.replace('№', '').replace(' ', '')]['status'] in ['Направлено на регистрацию в РТН', 'Зарегистрировано', 'Зарегистрировано и передано в цех']:
                    exceptions_list.append(f'Заключение со статусом "{consolidated_file_dict[zepb_reg_num.replace('№', '').replace(' ', '')]['status']}" не может быть направлено на регистрацию')

            if file_name[-4:].replace('_', '/').replace('-', '/') not in zepb_reg_num:
                exceptions_list.append(f'Номер заключения ({zepb_reg_num}) не соответствует имени файла')

            if len(zepb_name) >= 1000:
                exceptions_list.append(f'Наименование заключения длиннее 1000 символов (длина - {len(zepb_name)})')

            # составляем список номеров ОПО из файла для последующей проверки данных заявления
            opo_number_list = []
            for row in opo_address_file.tables[0].rows:
                opo_number_list.append(row.cells[1].text)

            if opo_number not in opo_number_list:
                exceptions_list.append(f'В заявлении указан номер ОПО, который отсутствует в файле {opo_address_directory}')

            # извлекаем данные по ОПО из специального файла:
            for i, row in enumerate(opo_address_file.tables[0].rows, start=0):
                for cell in row.cells:
                    if opo_number in cell.text:
                        index = i
                        break

            opo_address = opo_address_file.tables[0].rows[index].cells[3].text.strip()
            opo_class = opo_address_file.tables[0].rows[index].cells[2].text.strip()
            opo_name = opo_address_file.tables[0].rows[index].cells[0].text.strip()

            if any([opo_address=='', opo_class=='', opo_name=='']):
                exceptions_list.append(f'По ОПО №{opo_number} в файле {opo_address_directory} внесены не все данные (адрес, класс или наименование)')

            epb_obj_numbers_list = []
            for j in ["зав.№", "зав. №", "зав №", "зав№", "зав.", "зав. ",   # сначала ищем заводские номера,
                      "поз.№", "поз. №", "поз №", "поз№", "поз.", "поз. ",   # позиционные,
                      "рег.№", "рег. №", "рег №", "рег№", "рег.", "рег. ",   # регистрационные,
                      "тех.№", "тех. №", "тех №", "тех№", "тех.", "тех. ",   # технологические,
                      "№"]:                                                  # или просто символ номера
                if j in identification_number:
                    zepb_obj_id_num_full_row = identification_number.strip()
                    while '  ' in zepb_obj_id_num_full_row:  # Заменяем двойные пробелы на одинарные
                        zepb_obj_id_num_full_row = zepb_obj_id_num_full_row.replace('  ', ' ')
                    for k in range(zepb_obj_id_num_full_row.count(j)):
                        zepb_obj_id_num_idx_start = zepb_obj_id_num_full_row.find(j) + len(j)
                        # Ищем запятую, двоеточие, "поз", "рег", "тех" или "зав",
                        # принимаем в качестве разделителя тот символ, который нашелся раньше,
                        # и отфильтровывем отрицательные значения, если функция find не находит символ (возвращает "-1")
                        try:
                            zepb_obj_id_num_idx_end = min(filter(lambda val: val > 0,
                                                            [zepb_obj_id_num_full_row.find(",", zepb_obj_id_num_idx_start),
                                                             zepb_obj_id_num_full_row.find(";", zepb_obj_id_num_idx_start),
                                                             zepb_obj_id_num_full_row.find("поз", zepb_obj_id_num_idx_start),
                                                             zepb_obj_id_num_full_row.find("рег", zepb_obj_id_num_idx_start),
                                                             zepb_obj_id_num_full_row.find("тех", zepb_obj_id_num_idx_start),
                                                             zepb_obj_id_num_full_row.find("зав", zepb_obj_id_num_idx_start)]))
                        # Если символы вообще не будут найдены, то принимаем, что заводской номер стоит в конце строки:
                        except ValueError:
                            zepb_obj_id_num_idx_end = len(zepb_obj_id_num_full_row)
                        zepb_obj_id_num = zepb_obj_id_num_full_row[zepb_obj_id_num_idx_start:zepb_obj_id_num_idx_end].strip()  # Извлекаем номер из строки
                        zepb_obj_id_num = zepb_obj_id_num.replace('№', '')  # Удаляем символ "№"
                        zepb_obj_id_num = zepb_obj_id_num.replace(j, '')  # Удаляем текст типа номера (например, "зав.")
                        zepb_obj_id_num = zepb_obj_id_num.strip()

                        try:
                            epb_obj_numbers_list.append(zepb_obj_id_num)  # Добавляем найденный номер в список
                        except:
                            exceptions_list.append('Перепроверьте идентификационные номера '
                                            'объектов экспертизы (заводской, технологический и прочее) в заявлении')

                        # Следующую итерацию начинаем с позиции, где был найден предыдущий номер:
                        zepb_obj_id_num_full_row = zepb_obj_id_num_full_row[zepb_obj_id_num_idx_end:]
                        zepb_obj_id_num_type = j  # Вносим тип номера в переменную для дальнейшего выбора чекбокса

                    # если список epb_obj_numbers_list полностью состоит из 'б/н' или 'н/д', то очищаем его
                    if epb_obj_numbers_list == [epb_obj_numbers_list.remove('б/н') for _ in range(epb_obj_numbers_list.count('б/н'))] or epb_obj_numbers_list == [epb_obj_numbers_list.remove('н/д') for _ in range(epb_obj_numbers_list.count('н/д'))]:
                        epb_obj_numbers_list = epb_obj_numbers_list.clear()

                    if epb_obj_numbers_list != [] and epb_obj_numbers_list != None and epb_obj_numbers_list != '':  # Завершаем поиск номеров только когда в список epb_obj_numbers_list внесли хотя бы один номер
                        break

            epb_obj_numbers_list = ', '.join(epb_obj_numbers_list)  # Приводим список к строчному виду с разделителем
            epb_obj_numbers_list = epb_obj_numbers_list.strip()  # Удаляем начальный и конечный пробел

            if epb_obj_numbers_list == '':
                exceptions_list.append('Перепроверьте идентификационные номера объектов экспертизы (заводской, технологический и прочее) в заявлении')

            # если строка с номерами заканчивается на точку, то убираем её:
            if epb_obj_numbers_list[-1] == '.':
                epb_obj_numbers_list = epb_obj_numbers_list[:-1]

            if len(epb_obj_numbers_list) >= 255:
                exceptions_list.append(f'Строка с идентификационными номерами (зав.№/тех.№/поз.№ или другое) длиннее 255 символов (длина - {len(epb_obj_numbers_list)})')

            # если дата подписания задана текстом:
            if type(zepb_sign_date_full_row) is str:
                zepb_sign_date = re.search(r'\d.{4,11}202\d', zepb_sign_date_full_row)[0].lower().replace('  ', ' ')

                # Приводим дату к формату "ДД.ММ.ГГГГ" (меняем месяц с "ММММ" на "ММ"):
                for i in date_dict.keys():
                    old_value = i
                    new_value = date_dict[i]
                    # Для применения замены к переменной zepb_sign_date:
                    zepb_sign_date = zepb_sign_date.replace(old_value, new_value)

                # Преобразуем дату подписания в объект datetime:
                zepb_sign_date = datetime.strptime(zepb_sign_date, '%d.%m.%Y')
            # если дата подписания уже является объектом datetime, то ничего не меняем:
            elif type(zepb_sign_date_full_row) is datetime:
                zepb_sign_date = zepb_sign_date_full_row

            if type(zepb_exploperiod_full_row) is str:
                # Извлекаем дату из срока безопасной эксплуатации:
                if 'до ' in zepb_exploperiod_full_row:
                    zepb_exploperiod_idx_start = zepb_exploperiod_full_row.find('до ') + 3
                else:
                    zepb_exploperiod_idx_start = 0

                for j in ["г.", "включительно"]:
                    if j in zepb_exploperiod_full_row:
                        zepb_exploperiod_idx_end = zepb_exploperiod_full_row.find(j)
                        break
                    else:
                        zepb_exploperiod_idx_end = len(zepb_exploperiod_full_row)

                zepb_exploperiod = zepb_exploperiod_full_row[zepb_exploperiod_idx_start:zepb_exploperiod_idx_end].strip()

                # если строка заканчивается на точку, то убираем её:
                if zepb_exploperiod[-1] == '.':
                    zepb_exploperiod = zepb_exploperiod[:-1]

                # Приводим дату к формату "ДД.ММ.ГГГГ" (меняем месяц с "ММММ" на "ММ"):
                for month in date_dict.keys():
                    zepb_exploperiod = zepb_exploperiod.replace(month, date_dict[month])

                # Преобразуем дату в объект datetime:
                zepb_exploperiod = datetime.strptime(zepb_exploperiod, '%d.%m.%Y')

            elif type(zepb_exploperiod_full_row) is datetime:    # если срок эксплуатации уже является объектом datetime
                zepb_exploperiod = zepb_exploperiod_full_row

            # Если срок безопасной эксплуатации указан без фразы "включительно", то принимаем срок на день ранее:
            if "включительно" not in str(zepb_exploperiod_full_row) and exp_organization == 'Экспертная организация 1':
                zepb_exploperiod = zepb_exploperiod - timedelta(days=1)

            # Проверяем наличие ошибок и возвращаем exception, если...

            # номер заключения не найден в файле сводной
            if zepb_reg_num.replace('№', '').replace(' ', '') not in consolidated_file_dict.keys():
                exceptions_list.append(f'Заключение отсутствует в файле сводной ({consolidated_file_directory})')

            # указанный в заявлении номер не соответствует имени файла
            if file_name.replace('_', '/').replace('-', '/') not in zepb_reg_num and len(file_name) == 8:
                exceptions_list.append(f'Регистрационный номер, указанный в заявлении ({zepb_reg_num}) '
                                       f'не соответствует наименованию файла ({path_to_statement_file})')

            # наименование объекта экспертизы не соответствует наименованию заключения ЭПБ
            if (zepb_object_name.replace('«', '').replace('»', '')
                    .replace('техническое устройство', '').replace('Техническое устройство', '')
                    .replace('технические устройства', '').replace('Технические устройства', '')
                    .replace('сооружение', '').replace('Сооружение', '')
                    .replace('здание', '').replace('Здание', '')
                    .strip().lower()
                    not in zepb_name
                            .replace('«', '').replace('»', '')
                            .replace('техническое устройство', '').replace('Техническое устройство', '')
                            .replace('технические устройства', '').replace('Технические устройства', '')
                            .replace('сооружение', '').replace('Сооружение', '')
                            .replace('здание', '').replace('Здание', '')
                            .strip().lower()):
                exceptions_list.append('Наименование объекта экспертизы не сооответствует наименованию заключения ЭПБ '
                                'или наименование объекта экспертизы и его назначение не разделены словом '
                                '"предназначен"/"предназначена"/"предназначены"')

            if len(str(expert_organization_INN)) != 10:
                exceptions_list.append('В заявлении есть ошибка в номере ИНН экспертной организации')

            if len(str(expert_organization_OGRN)) != 13:
                exceptions_list.append('В заявлении есть ошибка в номере ОРГН экспертной организации')

            if len(str(customer_INN)) != 10:
                exceptions_list.append('В заявлении есть ошибка в номере ИНН организации-заказчика')

            if exceptions_list != []:
                raise Exception('|'.join(exceptions_list))

            # при отсутствии ошибок начинаем работу в браузере
            if only_check is not True:
                if file_check == input.split(separator)[0]:
                    click_on_pic('_internal/Screenshots/1_Browser.png')
                click_on_pic('_internal/Screenshots/2_Vedenie_reestra.png')
                while True:
                    try:
                        location = pyautogui.locateOnScreen('_internal/Screenshots/5_Nachat.png', confidence=confidence)
                        point = pyautogui.center(location)
                        x, y = point
                        pyautogui.click(x, y)
                        break
                    except pyautogui.ImageNotFoundException:  # Если есть черновик заявления
                        try:
                            location = pyautogui.locateOnScreen('_internal/Screenshots/3_Udalit.png', confidence=confidence)
                            point = pyautogui.center(location)
                            x, y = point
                            pyautogui.click(x, y)
                            time.sleep(8)
                            click_on_pic('_internal/Screenshots/4_Sozdat_Z.png')
                            click_on_pic('_internal/Screenshots/5_Nachat.png')
                            break
                        except pyautogui.ImageNotFoundException:
                            continue
                click_on_pic('_internal/Screenshots/6_Vnesti_v_reestr.png')
                click_on_pic('_internal/Screenshots/7_Nachat.png')
                click_on_pic('_internal/Screenshots/8_Pereyti_k_zayav.png')
                click_on_pic('_internal/Screenshots/9_Verno.png')
                time.sleep(1)
                click_on_pic('_internal/Screenshots/10_Verno.png')
                click_on_pic('_internal/Screenshots/11_Telefon.png')
                keyboard.press_and_release("ctrl+a")
                paste_info(telephone)
                click_on_pic('_internal/Screenshots/12_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/13_e_mail.png')
                keyboard.press_and_release("ctrl+a")
                paste_info(e_mail)
                click_on_pic('_internal/Screenshots/14_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/15_Polnoe_naimenovanie.png')
                paste_info(expert_organization_full_name)
                click_on_pic('_internal/Screenshots/16_INN.png')
                paste_info(expert_organization_INN)
                click_on_pic('_internal/Screenshots/17_OGRN.png')
                paste_info(expert_organization_OGRN)
                click_on_pic('_internal/Screenshots/18_Nomer_licenz.png')
                paste_info(license_number)
                pyautogui.scroll(-100)
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/19_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/20_Nomer_kval_udost.png')
                paste_info(expert_number_list[0])
                time.sleep(0.3)

                if len(expert_number_list) > 1:  # если есть второй эксперт
                    click_on_pic('_internal/Screenshots/21_Dobavit_eksp.png')
                    time.sleep(0.3)
                    keyboard.press_and_release("end")
                    time.sleep(0.3)
                    click_on_pic('_internal/Screenshots/21-1_Nomer_kval_udost.png')
                    paste_info(expert_number_list[1])
                if len(expert_number_list) > 2:  # если есть третий эксперт
                    click_on_pic('_internal/Screenshots/21_Dobavit_eksp.png')
                    time.sleep(0.3)
                    click_on_pic('_internal/Screenshots/21-1_Nomer_kval_udost.png')
                    paste_info(expert_number_list[2])
                    keyboard.press_and_release("end")
                    time.sleep(0.3)

                click_on_pic('_internal/Screenshots/22_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/22-1_Vyber_pol.png')
                keyboard.press_and_release("home")
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/23_Naimen_zakl.png')
                paste_info(zepb_name)
                pyautogui.scroll(-100)
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/24_Obekty_ekspert.png')

                if "ехническ" in zepb_name:
                    paste_info("Технические устройства")
                    time.sleep(0.2)
                    click_on_pic('_internal/Screenshots/25_Tekhicheskie_ustr.png')
                elif "здани" in zepb_name and "ехническ" not in zepb_name:
                    click_on_pic('_internal/Screenshots/25-1_Zdaniya_sooruzh.png')
                elif "сооружени" in zepb_name and "ехническ" not in zepb_name:
                    click_on_pic('_internal/Screenshots/25-1_Zdaniya_sooruzh.png')

                pyautogui.scroll(-100)
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/26_Naimen_obekta.png')
                paste_info(zepb_object_name)
                pyautogui.scroll(-100)
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/27_Naznachenie_obekta.png')
                paste_info(zepb_object_assignment)
                time.sleep(0.2)
                pyautogui.scroll(-500)
                time.sleep(0.2)

                if 'зав' in zepb_obj_id_num_type:
                    click_on_pic('_internal/Screenshots/28_Zavod_nomer.png')
                elif 'рег' in zepb_obj_id_num_type:
                    click_on_pic('_internal/Screenshots/29_Reg_nomer.png')
                elif 'тех' in zepb_obj_id_num_type:
                    click_on_pic('_internal/Screenshots/30_Tekh_nomer.png')
                elif 'поз' in zepb_obj_id_num_type:
                    click_on_pic('_internal/Screenshots/30_Tekh_nomer.png')
                else:
                    click_on_pic('_internal/Screenshots/28_Zavod_nomer.png')

                keyboard.press_and_release("end")
                time.sleep(0.3)
                click_on_pic('_internal/Screenshots/31_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/32_Pole_nomera.png')
                paste_info(epb_obj_numbers_list)
                click_on_pic('_internal/Screenshots/33_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/33-1_Ukazhite_dann.png')

                while True:
                    try:
                        location = pyautogui.locateOnScreen('_internal/Screenshots/34_Polnoe_naim.png', confidence=confidence)
                        point = pyautogui.center(location)
                        x, y = point
                        pyautogui.click(x, y)
                        confidence = 0.9
                        break
                    except pyautogui.ImageNotFoundException:
                        confidence -= 0.05
                        continue

                paste_info(customer_full_name)

                while True:
                    try:
                        location = pyautogui.locateOnScreen('_internal/Screenshots/35_INN.png', confidence=confidence)
                        point = pyautogui.center(location)
                        x, y = point
                        pyautogui.click(x, y)
                        confidence = 0.9
                        break
                    except pyautogui.ImageNotFoundException:
                        confidence -= 0.05
                        continue

                paste_info(customer_INN)
                click_on_pic('_internal/Screenshots/36_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/37_Naimenovanie_proiz_obekta.png')
                paste_info(opo_name)
                pyautogui.scroll(-100)
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/38_Klass_opas.png')

                while True:
                    try:
                        location = pyautogui.locateOnScreen('_internal/Screenshots/39_Klassy_opas_I.png', confidence=confidence)
                        break
                    except pyautogui.ImageNotFoundException:
                        time.sleep(0.7)
                        click_on_pic('_internal/Screenshots/37-1_Klass_opas.png')
                        click_on_pic('_internal/Screenshots/38_Klass_opas.png')
                        continue

                if opo_class == 'I':
                    click_on_pic('_internal/Screenshots/39_Klassy_opas_I.png')
                elif opo_class == 'II':
                    click_on_pic('_internal/Screenshots/40_Klassy_opas_II.png')
                elif opo_class == 'III':
                    click_on_pic('_internal/Screenshots/41_Klassy_opas_III.png')
                elif opo_class == 'IV':
                    click_on_pic('_internal/Screenshots/42_Klassy_opas_IV.png')

                click_on_pic('_internal/Screenshots/43_Reg_nomer_OPO.png')
                paste_info(opo_number)
                click_on_pic('_internal/Screenshots/44_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/45_Territoria.png')
                click_on_pic('_internal/Screenshots/46_Vybrat_iz_spravoch.png')
                paste_info(customer_region)
                click_on_pic('_internal/Screenshots/47_Kras_kray.png')
                pyautogui.scroll(-250)
                time.sleep(0.5)
                click_on_pic('_internal/Screenshots/48_Utochnite_mesto.png')
                paste_info(opo_address)
                time.sleep(0.3)
                pyautogui.move(500, 0)
                pyautogui.click()
                keyboard.press_and_release("end")
                time.sleep(0.3)
                click_on_pic('_internal/Screenshots/49_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/50_Vyvody.png')
                paste_info(' ')
                time.sleep(0.5)
                click_on_pic('_internal/Screenshots/51_Obekt_eksp_sootvet.png')
                click_on_pic('_internal/Screenshots/52_Reg_nomer_zakl.png')
                paste_info(zepb_reg_num)
                pyautogui.scroll(-80)
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/53_Data_podpis.png')
                paste_info(zepb_sign_date.strftime("%d.%m.%Y"))
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/54_Srok_bez_ekspl.png')
                paste_info(zepb_exploperiod.strftime("%d.%m.%Y"))
                time.sleep(0.5)
                click_on_pic('_internal/Screenshots/55_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/55-1_Vyberite_territ_organ.png')

                while True:
                    try:
                        location = pyautogui.locateOnScreen('_internal/Screenshots/56_Territ_organ.png', confidence=confidence)
                        point = pyautogui.center(location)
                        x, y = point
                        pyautogui.click(x, y)
                        confidence = 0.9
                        break
                    except pyautogui.ImageNotFoundException:
                        confidence -= 0.05
                        continue

                paste_info(rtn_region)
                time.sleep(1.5)
                keyboard.press_and_release("down")
                time.sleep(0.5)
                keyboard.press_and_release("enter")
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/58_Prodolzhit.png')
                statement_number = 'Не определено'
                statement_address = 'Не определено'

                pyautogui.scroll(-100)
                time.sleep(0.2)
                click_on_pic('_internal/Screenshots/59_Vybrat_fail.png')
                click_on_pic('_internal/Screenshots/59-1_Pole.png')
                time.sleep(0.7)
                paste_info(path_to_zepb)
                keyboard.press_and_release("enter")
                time.sleep(0.5)
                while True:
                    try:
                        location = pyautogui.locateOnScreen('_internal/Screenshots/59-1_Pole.png', confidence=confidence)
                        point = pyautogui.center(location)
                        x, y = point
                        pyautogui.click(x, y)
                        paste_info(path_to_zepb)
                        keyboard.press_and_release("enter")
                        click_on_pic('_internal/Screenshots/59-2_Zagruzh.png')
                        break
                    except pyautogui.ImageNotFoundException:  # Если есть черновик заявления
                        try:
                            location = pyautogui.locateOnScreen('_internal/Screenshots/59-2_Zagruzh.png', confidence=confidence)
                            point = pyautogui.center(location)
                            x, y = point
                            pyautogui.click(x, y)
                            break
                        except pyautogui.ImageNotFoundException:
                            continue
                keyboard.press_and_release("end")
                click_on_pic('_internal/Screenshots/60_Prodolzhit.png')
                click_on_pic('_internal/Screenshots/60-1_Podpishite.png')

                while True:
                    try:
                        location = pyautogui.locateOnScreen('_internal/Screenshots/61_Skachayte_fail.png', confidence=confidence)
                        point = pyautogui.center(location)
                        x, y = point
                        pyautogui.click(x, y)
                        confidence = 0.9
                        break
                    except pyautogui.ImageNotFoundException:
                        confidence -= 0.03
                        continue

                time.sleep(1.5)

                def working_with_pdf():

                    click_on_pic('_internal/Screenshots/62_Vyravnivanie.png')
                    click_on_pic('_internal/Screenshots/63_Podognat.png')
                    keyboard.press_and_release("end")
                    click_on_pic('_internal/Screenshots/64_Instrumenty.png')
                    click_on_pic('_internal/Screenshots/65_Sertificaty.png')
                    click_on_pic('_internal/Screenshots/66_Postavit_cifr_podp.png')
                    click_on_pic('_internal/Screenshots/66_Postavit_cifr_podp.png', y_offset=100, action='moveTo')

                    # Подписываем заявление
                    confidence = signing_statement_confidence
                    while True:
                        try:
                            location = pyautogui.locateOnScreen('_internal/Screenshots/66-1_Podpishite_zayav.png', confidence=confidence)
                            point = pyautogui.center(location)
                            x, y = point
                            pyautogui.moveTo(x + startpoint_x_deviation, y - startpoint_y_deviation)
                            pyautogui.dragTo(x + endpoint_x_deviation, y + endpoint_y_deviation, signing_speed, button='left')
                            confidence = 0.9
                            break
                        except pyautogui.ImageNotFoundException:
                            confidence -= signing_confidence_reduction_speed
                            continue

                    click_on_pic('_internal/Screenshots/67_Certificat_SDA.png')
                    click_on_pic('_internal/Screenshots/69_Ok.png')
                    click_on_pic('_internal/Screenshots/70_Prodolzhit.png')
                    time.sleep(0.7)
                    click_on_pic('_internal/Screenshots/71_Checkbox.png')
                    click_on_pic('_internal/Screenshots/72_Podpisat.png')
                    click_on_pic('_internal/Screenshots/73_Sohranit.png')
                    click_on_pic('_internal/Screenshots/74_Da.png')
                    time.sleep(0.5)
                    click_on_pic('_internal/Screenshots/75_Da.png')
                    time.sleep(2)

                # если в папке загрузок уже есть файл заявления с таким же именем, как у текущего, то удаляем его:
                if os.path.isfile(fr"{downloads_directory}/{file_name}.pdf") is True:
                    os.remove(fr"{downloads_directory}/{file_name}.pdf")

                # проверяем, появился ли файл заявления в папке загрузок, и переименовываем его после 0,5 секунд паузы
                while True:
                    if os.path.isfile(fr"{downloads_directory}/pdf.pdf") is True:
                        time.sleep(0.5)
                        while True:
                            try:
                                os.rename(fr"{downloads_directory}/pdf.pdf", fr"{downloads_directory}/{file_name}.pdf")
                                break
                            except:
                                message(f'Не удалось переименовать файл уведомления: {downloads_directory}/pdf.pdf".\n'
                                               'Проверьте, не открыт ли он и при необходимости закройте, после чего нажмите "Ок"')
                                continue
                        break
                    elif os.path.isfile(fr"{downloads_directory}/pdf.pdf") is False:
                        continue

                try:
                    # пробуем открыть файл заявления с помощью библиотеки subprocess
                    file = fr"{downloads_directory}/{file_name}.pdf"
                    prog = fr"{acrobat_reader_path}"
                    OpenIt = subprocess.Popen([prog, file])
                    working_with_pdf()
                    OpenIt.terminate()
                except FileNotFoundError:
                    # в случае ошибки открываем файл заявления с помощью библиотеки os
                    message('Возникла ошибка открытия файла pdf с помощью subprocess. Открываем файл с помощью os.'
                                   'Для продолжения нажмите Ok.')
                    os.startfile(fr"{downloads_directory}/{file_name}.pdf")
                    working_with_pdf()
                    keyboard.press_and_release("alt+F4")

                time.sleep(2)

                # если папка Archivation не существует, создаем её:
                if not os.path.isdir(Path(os.getcwd(), '_internal', 'Archivation')):
                    os.mkdir(Path(os.getcwd(), '_internal', 'Archivation'))

                # очищаем папку Archivation перед перемещением туда файлов
                directory = Path.cwd() / os.getcwd() / '_internal/Archivation'
                filesToRemove = [os.path.join(directory, file) for file in os.listdir(directory)]

                try:
                    for file in filesToRemove:
                        os.remove(file)
                except:
                    pass

                # перемещаем файл заявления в папку Archivation для последующего архивирования
                n = 0
                while True:
                    try:
                        shutil.move(src=fr"{downloads_directory}/{file_name}.pdf",
                                    dst=Path.cwd() / os.getcwd() / '_internal/Archivation')
                        break
                    except:
                        n += 1
                        if n >= 5:
                            message('Бот не может переместить файл из папки Downloads в папку Archivation.\n'
                                           'Переместите самостоятельно и нажмите Ок')
                            break
                        time.sleep(1)
                        continue

                time.sleep(1)

                # Архивируем файл заявления
                shutil.make_archive(format='zip',
                                    base_name=Path(fr'{downloads_directory}/{file_name}'),
                                    root_dir=Path.cwd() / os.getcwd() / '_internal/Archivation')

                time.sleep(0.5)

                # Перемещаем файл заявления в папку. Если там уже есть заявление с таким именем, сначала удаляем его:
                if os.path.isfile(fr"{statement_directory}/{file_name}.pdf") is True:
                    os.remove(fr"{statement_directory}/{file_name}.pdf")
                shutil.move(src=Path.cwd() / os.getcwd() / rf'_internal/Archivation/{file_name}.pdf',
                            dst=fr"{statement_directory}")

                click_on_pic('_internal/Screenshots/76_Vybrat_fail.png', confidence=0.8)
                click_on_pic('_internal/Screenshots/76-1_Pole.png')
                paste_info(rf"{downloads_directory}\{file_name}.zip")
                keyboard.press_and_release("enter")
                click_on_pic('_internal/Screenshots/77_Otpravit.png')
                click_on_pic('_internal/Screenshots/78_Lich_kabinet.png')
                click_on_pic('_internal/Screenshots/79_Pervoe_zayavlenie.png')
                click_on_pic('_internal/Screenshots/80_Naimenovanie_uslugi.png')
                keyboard.press_and_release("ctrl+l")
                time.sleep(0.5)
                keyboard.press_and_release("ctrl+insert")
                statement_address = pyperclip.paste()
                # Извлекаем номер заявления из pdf-файла
                reader = PdfReader(fr"{statement_directory}/{file_name}.pdf")
                page = reader.pages[0]
                statement_number = int(page.extract_text().strip()[30:40])

                consolidated_file_cell_coordinate = None
                consolidated_file_cell = None
                consolidated_file_cell_coordinate = consolidated_file_dict[zepb_reg_num.replace('№', '').replace(' ', '')]['coordinate']
                consolidated_file_cell = ws_cf[consolidated_file_cell_coordinate]

                consolidated_file_cell.offset(column=3).value = 'Направлено на регистрацию в РТН'
                consolidated_file_cell.offset(column=4).value = datetime.now()
                consolidated_file_cell.offset(column=9).value = statement_address
                consolidated_file_cell.offset(column=8).value = statement_number

                while True:
                    try:
                        wb_cf.save(consolidated_file_directory)
                        break
                    except PermissionError:
                        message(f'Не удалось сохранить файл {consolidated_file_directory}.'
                                       'Проверьте, не открыт ли он и при необходимости закройте, после чего нажмите Ок')
                        continue

                # Удаляем архивный файл с заявлением из папки загрузок
                while True:
                    try:
                        os.remove(fr"{downloads_directory}/{file_name}.zip")
                        break
                    except:
                        continue

            processed_files_list.append(file_name)
            if only_check is True:
                logging.info(f'Файл {file_name} не содержит ошибок.')
            else:
                logging.info(f'Файл {file_name} успешно обработан. '
                             f'Экспертная организация - {exp_organization}, управление - {rtn_region}, '
                             f'номер - {statement_number}, URL - {statement_address}.')

        except Exception as exc:
            error_files_list.append(file_name)
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logging.error(f'Возникли ошибки в файле {file_name} ({exp_organization}): {exc}. Линия: {exc_tb.tb_lineno}.')
            pass

    wb_cf.close()

    with open('ZEPB.log', 'a') as file:
        file.write('\n')

    processed_files = '\n'.join(processed_files_list)
    error_files = '\n'.join(error_files_list)

    if only_check is not True:
        if len(error_files_list) == 0:
            message(f'Готово. Обработаны заключения ({len(processed_files_list)}):\n{processed_files}.')
        elif len(error_files_list) > 0 and len(processed_files_list) > 0:
            message(f'Готово. Обработаны заключения ({len(processed_files_list)}):\n{processed_files}.'
                           f'\nВозникли ошибки у заключений ({len(error_files_list)}):\n{error_files}.'
                           f'\nПодробности в файле {Path(os.getcwd(), 'ZEPB.log')}')
            os.startfile('ZEPB.log')
        elif len(error_files_list) > 0 and len(processed_files_list) == 0:
            message(f'Заключения не были обработаны - возникли ошибки у заключений ({len(error_files_list)}):\n{error_files}.'
                           f'\nПодробности в файле {Path(os.getcwd(), 'ZEPB.log')}')
            os.startfile('ZEPB.log')
    else:
        if len(error_files_list) == 0:
            message(f'Ни одно из указанных заявлений не содержит ошибок ({len(processed_files_list)}):\n{processed_files}.')
        elif len(error_files_list) > 0 and len(processed_files_list) > 0:
            message(f'Не содержат ошибок ({len(processed_files_list)}):\n{processed_files}.'
                           f'\nЕсть ошибки в ({len(error_files_list)}):\n{error_files}.'
                           f'\nПодробности в файле {Path(os.getcwd(), 'ZEPB.log')}')
            os.startfile('ZEPB.log')
        elif len(error_files_list) > 0 and len(processed_files_list) == 0:
            message(f'Во всех заявлениях есть ошибки ({len(error_files_list)}):\n{error_files}.'
                           f'\nПодробности в файле {Path(os.getcwd(), 'ZEPB.log')}')
            os.startfile('ZEPB.log')